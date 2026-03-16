"""
DOCX document chunker for regulation search.

Implements the chunking strategy from the specification:
1. Extract paragraphs, tables, and headings from docx
2. Exclude TOC paragraphs
3. Detect headings by style or bold formatting
4. Convert 1x1 tables to text paragraphs
5. Convert other tables to markdown with merged cells
6. Tables in markdown become 1:1 fragments
7. Remove service marks like [Новое]
8. Merge small paragraphs within same section
9. Split large paragraphs by sentences
10. Medium paragraphs become 1:1 fragments
11. Prepend nearest heading to each fragment
"""

import re
from dataclasses import dataclass, field
from docx import Document
from docx.table import Table
from docx.oxml.ns import qn


@dataclass
class Fragment:
    text: str
    heading: str
    doc_name: str
    fragment_type: str  # "text", "table", "heading"
    position: int = 0


def _is_toc_style(style_name: str) -> bool:
    """Check if paragraph style is TOC (table of contents)."""
    if not style_name:
        return False
    return style_name.lower().startswith("toc") or "содержание" in style_name.lower()


def _is_heading_style(style_name: str) -> bool:
    """Check if paragraph style is a heading."""
    if not style_name:
        return False
    name = style_name.lower()
    return (
        name.startswith("heading")
        or name.startswith("заголовок")
        or "title" in name
        or "subtitle" in name
    )


def _is_all_bold(paragraph) -> bool:
    """Check if the entire paragraph is bold text."""
    runs = [r for r in paragraph.runs if r.text.strip()]
    if not runs:
        return False
    return all(r.bold for r in runs)


def _is_heading(paragraph) -> bool:
    """Detect headings by style or bold formatting."""
    style_name = paragraph.style.name if paragraph.style else ""
    if _is_heading_style(style_name):
        return True
    text = paragraph.text.strip()
    if text and _is_all_bold(paragraph) and len(text) < 200:
        return True
    return False


def _clean_text(text: str) -> str:
    """Remove service marks and clean text."""
    text = re.sub(r'\[(?:Н|н)овое\]', '', text, flags=re.IGNORECASE)
    text = re.sub(r'\u200b', '', text)  # zero-width space
    text = re.sub(r'\xa0', ' ', text)  # non-breaking space
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def _get_merged_cell_text(table: Table) -> list[list[str]]:
    """
    Extract table cells handling merged cells.
    Returns a 2D list of cell texts with merged cell markers.
    """
    rows = []
    grid = table._tbl
    tr_elements = grid.findall(qn('w:tr'))

    for tr in tr_elements:
        row_cells = []
        tc_elements = tr.findall(qn('w:tc'))
        for tc in tc_elements:
            text = ""
            for p in tc.findall(qn('w:p')):
                p_text = ""
                for r in p.findall(qn('w:r')):
                    for t in r.findall(qn('w:t')):
                        if t.text:
                            p_text += t.text
                if p_text:
                    text += p_text + " "
            text = _clean_text(text)

            # Check colspan (gridSpan)
            tc_pr = tc.find(qn('w:tcPr'))
            colspan = 1
            if tc_pr is not None:
                grid_span = tc_pr.find(qn('w:gridSpan'))
                if grid_span is not None:
                    colspan = int(grid_span.get(qn('w:val'), 1))

            row_cells.append(text)
            # Add empty cells for merged columns
            for _ in range(colspan - 1):
                row_cells.append(text)  # duplicate content for merged cells

        rows.append(row_cells)
    return rows


def _table_to_markdown(table: Table) -> str:
    """Convert a docx table to markdown format with merged cells handled."""
    rows = _get_merged_cell_text(table)
    if not rows:
        return ""

    # Normalize column count
    max_cols = max(len(r) for r in rows)
    for r in rows:
        while len(r) < max_cols:
            r.append("")

    lines = []
    for i, row in enumerate(rows):
        line = "| " + " | ".join(cell if cell else " " for cell in row) + " |"
        lines.append(line)
        if i == 0:
            lines.append("| " + " | ".join(["---"] * max_cols) + " |")

    return "\n".join(lines)


def _split_sentences_simple(text: str) -> list[str]:
    """
    Simple sentence splitter (Stage 1 - without Natasha).
    Splits by sentence-ending punctuation followed by space/newline.
    """
    # Split on sentence boundaries: .!? followed by space and uppercase or end
    sentences = re.split(r'(?<=[.!?])\s+(?=[А-ЯA-Z0-9])', text)
    return [s.strip() for s in sentences if s.strip()]


def _merge_small_chunks(texts: list[str], min_chars: int, max_chars: int) -> list[str]:
    """Merge small text chunks that are in the same section."""
    if not texts:
        return []

    merged = []
    current = texts[0]

    for text in texts[1:]:
        if len(current) < min_chars and len(current) + len(text) + 1 <= max_chars:
            current = current + "\n" + text
        else:
            merged.append(current)
            current = text

    merged.append(current)
    return merged


def _split_large_text(text: str, max_chars: int) -> list[str]:
    """Split large text into smaller fragments by sentences."""
    if len(text) <= max_chars:
        return [text]

    sentences = _split_sentences_simple(text)
    if len(sentences) <= 1:
        return [text]

    chunks = []
    current = sentences[0]

    for sentence in sentences[1:]:
        if len(current) + len(sentence) + 1 <= max_chars:
            current = current + " " + sentence
        else:
            chunks.append(current)
            current = sentence

    if current:
        chunks.append(current)

    return chunks


def chunk_document(file_path: str, min_chars: int = 100, max_chars: int = 1500) -> list[Fragment]:
    """
    Process a DOCX document and return a list of text fragments.

    Args:
        file_path: Path to the .docx file
        min_chars: Minimum chunk size in characters
        max_chars: Maximum chunk size in characters

    Returns:
        List of Fragment objects
    """
    import os
    doc_name = os.path.splitext(os.path.basename(file_path))[0]
    doc = Document(file_path)

    # First pass: collect all elements in order (paragraphs + tables)
    elements = []
    body = doc.element.body

    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

        if tag == 'p':
            # Find matching paragraph object
            for para in doc.paragraphs:
                if para._element is child:
                    elements.append(('paragraph', para))
                    break

        elif tag == 'tbl':
            # Find matching table object
            for tbl in doc.tables:
                if tbl._tbl is child:
                    elements.append(('table', tbl))
                    break

    # Second pass: process elements into fragments
    raw_fragments = []
    current_heading = ""
    position = 0

    for elem_type, elem in elements:
        if elem_type == 'paragraph':
            style_name = elem.style.name if elem.style else ""

            # Skip TOC paragraphs
            if _is_toc_style(style_name):
                continue

            text = _clean_text(elem.text)
            if not text:
                continue

            # Check if heading
            if _is_heading(elem):
                current_heading = text
                continue

            raw_fragments.append({
                'text': text,
                'heading': current_heading,
                'type': 'text',
            })

        elif elem_type == 'table':
            rows = len(elem.rows)
            cols = len(elem.columns)

            if rows == 1 and cols == 1:
                # 1x1 table -> text paragraph
                text = _clean_text(elem.cell(0, 0).text)
                if text:
                    raw_fragments.append({
                        'text': text,
                        'heading': current_heading,
                        'type': 'text',
                    })
            elif cols == 1:
                # Single-column table (Nx1) -> treat each row as text paragraph
                for row in elem.rows:
                    text = _clean_text(row.cells[0].text)
                    if text:
                        # Check if this cell looks like a heading
                        if len(text) < 200 and not any(c in text for c in '.;:,'):
                            current_heading = text
                        else:
                            raw_fragments.append({
                                'text': text,
                                'heading': current_heading,
                                'type': 'text',
                            })
            else:
                # Multi-cell table -> markdown
                md = _table_to_markdown(elem)
                if md:
                    raw_fragments.append({
                        'text': md,
                        'heading': current_heading,
                        'type': 'table',
                    })

    # Third pass: group text fragments by heading, merge small, split large
    result_fragments = []
    position = 0

    # Group consecutive text fragments with same heading
    groups = []
    current_group = []
    current_group_heading = None

    for frag in raw_fragments:
        if frag['type'] == 'table':
            # Tables are 1:1 fragments, flush current group first
            if current_group:
                groups.append((current_group_heading, current_group))
                current_group = []
            result_fragments.append(Fragment(
                text=frag['text'],
                heading=frag['heading'],
                doc_name=doc_name,
                fragment_type='table',
                position=position,
            ))
            position += 1
            current_group_heading = frag['heading']
        else:
            if frag['heading'] != current_group_heading and current_group:
                groups.append((current_group_heading, current_group))
                current_group = []
            current_group_heading = frag['heading']
            current_group.append(frag['text'])

    if current_group:
        groups.append((current_group_heading, current_group))

    # Process text groups
    for heading, texts in groups:
        # Merge small paragraphs
        merged = _merge_small_chunks(texts, min_chars, max_chars)

        for text in merged:
            # Split large paragraphs
            splits = _split_large_text(text, max_chars)

            for chunk in splits:
                if chunk.strip():
                    result_fragments.append(Fragment(
                        text=chunk,
                        heading=heading or "",
                        doc_name=doc_name,
                        fragment_type='text',
                        position=position,
                    ))
                    position += 1

    return result_fragments


def fragments_to_search_text(fragments: list[Fragment]) -> list[dict]:
    """
    Convert fragments to search-ready format with heading prepended.
    Returns list of dicts with 'text' (for embedding) and 'metadata'.
    """
    results = []
    for i, frag in enumerate(fragments):
        # Prepend heading to fragment text for better search context
        search_text = frag.text
        if frag.heading:
            search_text = f"{frag.heading}\n{frag.text}"

        results.append({
            'id': f"{frag.doc_name}_{i}",
            'text': search_text,
            'metadata': {
                'doc_name': frag.doc_name,
                'heading': frag.heading,
                'fragment_type': frag.fragment_type,
                'position': frag.position,
                'original_text': frag.text,
            }
        })

    return results
